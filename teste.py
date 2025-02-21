import os
from docx import Document
import tkinter as tk
from tkinter import messagebox, ttk
import locale
from datetime import datetime

def obter_data_formatada():
    try:
        locale.setlocale(locale.LC_TIME, "pt_BR.utf8")
    except locale.Error:
        locale.setlocale(locale.LC_TIME, "Portuguese_Brazil.1252")  # Windows
    return datetime.today().strftime("%d de %B de %Y")

# Atualiza a data automaticamente se o Checkbutton for marcado
def atualizar_data():
    if var1.get():  # Se estiver marcado
        entry_data.delete(0, tk.END)
        entry_data.insert(0, obter_data_formatada())
    else:
        entry_data.delete(0, tk.END)

# Verifica se todos os campos estão preenchidos antes de gerar os documentos
def verificar_campos():
    for entry in entries:
        if isinstance(entry, tk.Entry) and entry.get().strip() == "":
            messagebox.showwarning("Campo vazio", "Por favor, preencha todos os campos.")
            return False
    return True

# Geração dos documentos
def gerar_documentos():
    if not verificar_campos():
        return

    nome_completo = entry_nome.get()
    cpf = entry_cpf.get()
    rg = entry_rg.get()
    sexoF = combo_sexo.get() == "Feminino"
    estado_civil = combo_estado_civil.get()
    rua_e_numero = entry_rua.get()
    bairro = entry_bairro.get()
    data = entry_data.get()
    data_filiacao = entry_data_filiacao.get()
    nacionalidade = "BRASILEIRA" if sexoF else "BRASILEIRO"

    estado_civil_ajustado = {
        "Solteiro": "Solteira" if sexoF else "Solteiro",
        "Casado": "Casada" if sexoF else "Casado",
        "União estável": "União estável"
    }.get(estado_civil, estado_civil)

    doc = Document("./assets/residencia.docx")
    doc2 = Document("./assets/filiacao.docx")

    # Substituir textos no documento
    for para in doc.paragraphs:
        para.text = para.text.replace("FULANO BARBOSA", nome_completo)
        para.text = para.text.replace("BRASILEIRO", nacionalidade)
        para.text = para.text.replace("SOLTEIRO", estado_civil_ajustado)
        para.text = para.text.replace("000.000.000-00", cpf)
        para.text = para.text.replace("0000000-1", rg)
        para.text = para.text.replace("RUA SÃO BRAZ", rua_e_numero)
        para.text = para.text.replace("SUBSTAÇÃO", bairro)
        para.text = para.text.replace("28 de abril de 2024", data)

    for para in doc2.paragraphs:
        para.text = para.text.replace("MARIA FULANO", nome_completo)
        para.text = para.text.replace("000.000.000-00", cpf)
        para.text = para.text.replace("0000000-0", rg)
        para.text = para.text.replace("AV COELHO NETO", rua_e_numero)
        para.text = para.text.replace("CENTRO", bairro)
        para.text = para.text.replace("22 de agosto de 2024", data)

    for tabela in doc2.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                if "20/08/2015" in celula.text:
                    celula.text = "Data de Filiação:\n" + data_filiacao

    nova_pasta = os.path.join("declaracoes_geradas", nome_completo)
    os.makedirs(nova_pasta, exist_ok=True)

    doc.save(os.path.join(nova_pasta, "declaracao_residencia.docx"))
    doc2.save(os.path.join(nova_pasta, "declaracao_filiacao.docx"))

    messagebox.showinfo("Sucesso", f"Declarações de {nome_completo} criadas com sucesso!")

# Fecha a janela
def fechar_janela():
    root.destroy()

root = tk.Tk()
root.title("Gerador de Declarações")
root.config(padx=30, pady=30)

tk.Label(root, text="Gerador de Declarações", font=("Arial", 16, "bold")).grid(row=0, column=0, columnspan=2, pady=2)
tk.Label(root, text="Serão geradas declarações de FILIAÇÃO E RESIDÊNCIA em .docx", font=("Arial", 10, "bold"), fg="red").grid(row=1, column=0, columnspan=2, pady=0)

labels = ["Nome Completo", "CPF (com ou sem  .  e  -):", "RG (somente números)", "Estado civil:", "Sexo:", 'Rua e número ("Rua X, 123")', "Bairro", "Data de Filiação (DD/MM/AAAA)",'Data do Documento (ex: "5 de maio de 20XX")']
entries = []

for i, label in enumerate(labels):
    tk.Label(root, text=label).grid(row=i+2, column=0, padx=5, pady=5)
    if label == "Estado civil:":
        combo_estado_civil = ttk.Combobox(root, values=["Solteiro", "União estável", "Casado"], state="readonly", width=57)
        combo_estado_civil.grid(row=i+2, column=1, padx=5, pady=5)
        combo_estado_civil.current(0)
    elif label == "Sexo:":
        combo_sexo = ttk.Combobox(root, values=["Masculino", "Feminino"], state="readonly", width=57)
        combo_sexo.grid(row=i+2, column=1, padx=5, pady=5)
        combo_sexo.current(0)
    else:
        entry = tk.Entry(root, width=60)
        entry.grid(row=i+2, column=1, padx=5, pady=5)
        entries.append(entry)

entry_nome, entry_cpf, entry_rg, entry_rua, entry_bairro, entry_data_filiacao, entry_data = entries

var1 = tk.BooleanVar()
chk1 = tk.Checkbutton(root, text="Usar data de hoje", variable=var1, command=atualizar_data)
chk1.grid(row=10, column=1, padx=(250,0), pady=5)

btn_gerar = tk.Button(root, text="Gerar Documentos", bg="azure3", command=gerar_documentos)
btn_gerar.grid(row=13, column=0, columnspan=2, pady=10)

btn_fechar = tk.Button(root, text="Sair", command=fechar_janela, bg='red', fg='white', width=5)
btn_fechar.grid(row=13, column=1, columnspan=2, pady=10)

root.mainloop()

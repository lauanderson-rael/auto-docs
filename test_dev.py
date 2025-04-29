import os
from docx import Document
import tkinter as tk
from tkinter import messagebox , ttk
import locale
from datetime import datetime
from tkinter import filedialog, messagebox
from fillpdf import fillpdfs

def gerar_carteira(nome, data_nasc, data_emissao, pai, mae, cpf, rg, end, bairro, n_ficha, natu, estado_cv):

    try:
        campos_carteira = list(fillpdfs.get_form_fields("./assets/CARTEIRA-V4.pdf").keys())
        data_dict1 = {
            campos_carteira[0]: pai,
            campos_carteira[1]: n_ficha,
            campos_carteira[2]: nome,
            campos_carteira[3]: data_emissao,#
            campos_carteira[4]: mae,
            campos_carteira[5]: end,
            campos_carteira[6]: bairro,
            campos_carteira[7]: natu,
            campos_carteira[8]: data_nasc,
            campos_carteira[9]: estado_cv,
            campos_carteira[10]: rg,
            campos_carteira[11]: cpf,
            campos_carteira[12]: n_ficha,
            campos_carteira[13]: data_emissao,
        }

        output_dir = f"./declaracoes_geradas/{nome}"
        os.makedirs(output_dir, exist_ok=True)
        fillpdfs.write_fillable_pdf('./assets/CARTEIRA-V4.pdf', f'{output_dir}/carteira_de_socio.pdf', data_dict1)
        print("Carteira de socio criada com sucesso!")
    except Exception as e:
        messagebox.showerror("Erro", f"Ocorreu um erro: {e}")
#

def obter_data_formatada(modo="texto"):
    try:
        locale.setlocale(locale.LC_TIME, "pt_BR.utf8")
    except locale.Error:
        # Para Windows, caso "pt_BR.utf8" não funcione
        locale.setlocale(locale.LC_TIME, "Portuguese_Brazil.1252")

    if modo == "numerico":
        return datetime.today().strftime("%d/%m/%Y")
    else:  # modo == "texto"
        return datetime.today().strftime("%d de %B de %Y")

def gerar_documentos():
    if not verificar_campos():
        return

    nome_completo = entry_nome.get()
    cpf = entry_cpf.get()
    rg = entry_rg.get()
    sexoF = var_sexo.get() == 1
    rua_e_numero = entry_rua.get()
    bairro = entry_bairro.get()
    data = entry_data.get()
    data_filiacao = entry_data_filiacao.get()
    nacionalidade = 'BRASILEIRA' if sexoF else 'BRASILEIRO'

    estado_civil_base = entry_estado_civil.get()
    if estado_civil_base == "Solteiro(a)":
        estado_civil = "Solteira" if var_sexo.get() == 1 else "Solteiro"
    elif estado_civil_base == "Casado(a)":
        estado_civil = "Casada" if var_sexo.get() == 1 else "Casado"
    else:
        estado_civil = "União estável"

    # Carregar os documentos
    doc = Document("./assets/residencia.docx")
    doc2 = Document("./assets/filiacao.docx")

    # Textos a serem substituídos
    texto_informacoes = 'Na falta de documentos próprios, aptos que comprovem minha residência e domicílio, eu, FULANO BARBOSA, Nacionalidade: BRASILEIRO, estado Civil: SOLTEIRO, Profissão: PESCADOR(A), inscrito(a) no Cadastro de Pessoas Físicas (CPF) sob o nº 000.000.000-00, portador(a) da Carteira de Identidade (RG) nº 0000000-1, declaro ser residente e domiciliado (a) no endereço:  RUA SÃO BRAZ  Bairro: SUBSTAÇÃO'
    texto_data = 'Coelho Neto, 28 de abril de 2024'

    texto_informacoes2 = "Eu, MARIA FULANO, CPF:  000.000.000-00, RG: 0000000-0, residente no endereço completo: AV COELHO NETO, BAIRRO: CENTRO, declaro ser filiado à Entidade abaixo especificada:"
    texto_data2 = "Coelho Neto, 22 de agosto de 2024"

    # Percorrendo residencia
    for para in doc.paragraphs:
        if texto_informacoes in para.text:
            para.text = para.text.replace(texto_informacoes, f'Na falta de documentos próprios, aptos que comprovem minha residência e domicílio, eu, {nome_completo}, Nacionalidade: {nacionalidade}, estado Civil: {estado_civil}, Profissão: PESCADOR(A), inscrito(a) no Cadastro de Pessoas Físicas (CPF) sob o nº {cpf}, portador(a) da Carteira de Identidade (RG) nº {rg}, declaro ser residente e domiciliado (a) no endereço: {rua_e_numero}, Bairro: {bairro}')
        if texto_data in para.text:
            para.text = para.text.replace(texto_data, f'Coelho Neto, {data}')

    # Percorrendo filiacao
    for para in doc2.paragraphs:
        if texto_informacoes2 in para.text:
            para.text = para.text.replace(texto_informacoes2, f'Eu, {nome_completo}, CPF: {cpf}, RG: {rg}, residente no endereço completo: {rua_e_numero}, Bairro: {bairro}, declaro ser filiado à Entidade abaixo especificada:')
        if texto_data2 in para.text:
            para.text = para.text.replace(texto_data2, f'Coelho Neto, {data}')

    for tabela in doc2.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                if '20/08/2015' in celula.text:
                    celula.text= '' # Remove o texto original
                    paragrafo = celula.paragraphs[0]
                    paragrafo.add_run("Data de Filiação:\n")
                    paragrafo.add_run(data_filiacao).bold = True

    # salvando
    nova_pasta = os.path.join("declaracoes_geradas",nome_completo)
    os.makedirs(nova_pasta, exist_ok=True)

    caminho_arquivo = os.path.join(nova_pasta, "declaracao_residencia.docx")
    caminho_arquivo2 = os.path.join(nova_pasta, "declaracao_filiacao.docx")
    doc.save(caminho_arquivo)
    doc2.save(caminho_arquivo2)
    print("Declaracao criadas com sucesso!")

    gerar_carteira(
    nome_completo,
    entry_data_nasc.get(),
    data_filiacao,
    entry_pai.get(),
    entry_mae.get(),
    cpf,
    rg,
    rua_e_numero,
    bairro,
    entry_num_ficha.get(),
    entry_natu.get(),  # Pode ser fixo ou adicionado como campo também
    entry_estado_civil.get()
    )
    
    messagebox.showinfo("Sucesso", f"Declaração de residência, filiação e Carteira de Socio de {nome_completo} criadas com sucesso!")


def verificar_campos():
    for entry in entries:
        if entry.get().strip() == "":
            messagebox.showwarning("Campo vazio", "Por favor, preencha todos os campos.")
            return False
    return True


def fechar_janela():
    print("fechar")
    root.destroy()


# INTERFACE GRAFICA
root = tk.Tk()
root.title("Gerador de Declarações")
root.config(padx=30, pady=30)

# título H1
tk.Label(root, text="Gerador de Declarações", font=("Arial", 16, "bold")).grid(row=0, column=0, columnspan=2, pady=2)
tk.Label(root, text="Serão geradas as declarações de FILIAÇÃO, RESIDÊNCIA e CARTEIRA DE SÓCIO em .docx", font=("Arial", 10, "bold"), fg="red").grid(row=1, column=0, columnspan=2, pady=0)

# Labels e campos
labels = [
    "Nome Completo", "CPF:", "RG", "Estado civil:",
    'Rua e número ("Rua X, 123")', "Bairro",
    "Data de Filiação (DD/MM/AAAA)", 'Data do documento ("5 de maio de 20XX")',
    "Data de nascimento (DD/MM/AAAA)", "Naturalidade",
    "Nome do pai", "Nome da mãe", "Número da ficha"
]

entries = []

for i, label in enumerate(labels):
    # Labels verdes a partir do índice 8
    cor = "green" if i >= 8 else "black"
    peso = "bold" if i >= 8 else ' '
    tk.Label(root, text=label, fg=cor,  font=("default", 10, peso)).grid(row=i+2, column=0, padx=5, pady=5, sticky='w')

    if label == "Estado civil:":
        entry_estado_civil = ttk.Combobox(root, values=["Solteiro(a)", "Casado(a)", "União estável"], state="readonly", width=57)
        entry_estado_civil.grid(row=i+2, column=1, padx=5, pady=5)
        entry_estado_civil.current(0)
    else:
        entry = tk.Entry(root, width=60)
        entry.grid(row=i+2, column=1, padx=5, pady=5)
        entries.append(entry)

(entry_nome, entry_cpf, entry_rg, entry_rua, entry_bairro,
 entry_data_filiacao, entry_data, entry_data_nasc, entry_natu,
 entry_pai, entry_mae, entry_num_ficha) = entries

# valores padrao
entry_natu.insert(0,'Coelho Neto-MA')
entry_data_filiacao.insert(0, obter_data_formatada("numerico"))
entry_data.insert(0, obter_data_formatada())


# Sexo (Radio Buttons)
var_sexo = tk.IntVar()
frame_sexo = tk.Frame(root)
frame_sexo.grid(row=len(labels)+3, column=0, columnspan=2, pady=5)
tk.Radiobutton(frame_sexo, text="Sexo masculino", variable=var_sexo, value=0).pack(side="left", padx=10)
tk.Radiobutton(frame_sexo, text="Sexo feminino", variable=var_sexo, value=1).pack(side="left", padx=10)

# Botões
btn_frame = tk.Frame(root)
btn_frame.grid(row=len(labels)+4, column=0, columnspan=2, pady=10)
btn_gerar = tk.Button(btn_frame, text="Gerar Documentos", bg="azure3", command=gerar_documentos, width=20)
btn_gerar.pack(side="left", padx=10)
btn_fechar = tk.Button(btn_frame, text="Sair", command=fechar_janela, bg='red', fg='white', width=10)
btn_fechar.pack(side="left", padx=10)

root.mainloop()

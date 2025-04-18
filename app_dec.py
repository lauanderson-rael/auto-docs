import os
from docx import Document
import tkinter as tk
from tkinter import messagebox

import locale
from datetime import datetime

from tkinter import messagebox, ttk
#novo
def obter_data_formatada():
    try:
        locale.setlocale(locale.LC_TIME, "pt_BR.utf8")
    except locale.Error:
        # Para Windows, caso "pt_BR.utf8" não funcione
        locale.setlocale(locale.LC_TIME, "Portuguese_Brazil.1252")
    return datetime.today().strftime("%d de %B de %Y")
#novo

def gerar_documentos():
    if not verificar_campos():
        return

    # Obter os valores
    nome_completo = entry_nome.get()
    cpf = entry_cpf.get()
    rg = entry_rg.get()
    sexoF = var_sexo.get() == 1
    estado_civil = entry_estado_civil.get()
    rua_e_numero = entry_rua.get()
    bairro = entry_bairro.get()
    data = entry_data.get()
    data_filiacao = entry_data_filiacao.get()
    nacionalidade = 'BRASILEIRA' if sexoF else 'BRASILEIRO'

    # Carregar os documentos
    doc = Document("./assets/residencia.docx")
    doc2 = Document("./assets/filiacao.docx")

    # Textos a serem substituídos
    texto_informacoes = 'Na falta de documentos próprios, aptos que comprovem minha residência e domicílio, eu, FULANO BARBOSA, Nacionalidade: BRASILEIRO, estado Civil: SOLTEIRO, Profissão: PESCADOR(A), inscrito(a) no Cadastro de Pessoas Físicas (CPF) sob o nº 000.000.000-00, portador(a) da Carteira de Identidade (RG) nº 0000000-1, declaro ser residente e domiciliado (a) no endereço:  RUA SÃO BRAZ  Bairro: SUBSTAÇÃO'
    texto_data = 'Coelho Neto, 28 de abril de 2024'

    texto_informacoes2 = "Eu, MARIA FULANO, CPF:  000.000.000-00, RG: 0000000-0, residente no endereço completo: AV COELHO NETO, BAIRRO: CENTRO, declaro ser filiado à Entidade abaixo especificada:"
    texto_data2 = "Coelho Neto, 22 de agosto de 2024"

    # Percorrendo residencia
    for para in doc.paragraphs:
        if texto_informacoes in para.text:
            para.text = para.text.replace(texto_informacoes, f'Na falta de documentos próprios, aptos que comprovem minha residência e domicílio, eu, {nome_completo}, Nacionalidade: {nacionalidade}, estado Civil: {estado_civil}, Profissão: PESCADOR(A), inscrito(a) no Cadastro de Pessoas Físicas (CPF) sob o nº {cpf}, portador(a) da Carteira de Identidade (RG) nº {rg}, declaro ser residente e domiciliado (a) no endereço: {rua_e_numero}, Bairro: {bairro}')
        if texto_data in para.text:
            para.text = para.text.replace(texto_data, f'Coelho Neto, {data}')

    # Percorrendo filiacao
    for para in doc2.paragraphs:
        if texto_informacoes2 in para.text:
            para.text = para.text.replace(texto_informacoes2, f'Eu, {nome_completo}, CPF: {cpf}, RG: {rg}, residente no endereço completo: {rua_e_numero}, Bairro: {bairro}, declaro ser filiado à Entidade abaixo especificada:')
        if texto_data2 in para.text:
            para.text = para.text.replace(texto_data2, f'Coelho Neto, {data}')

    for tabela in doc2.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                if '20/08/2015' in celula.text:
                    celula.text= '' # Remove o texto original
                    paragrafo = celula.paragraphs[0]
                    paragrafo.add_run("Data de Filiação:\n")
                    paragrafo.add_run(data_filiacao).bold = True

    # salvando
    nova_pasta = os.path.join("declaracoes_geradas",nome_completo)
    os.makedirs(nova_pasta, exist_ok=True)

    caminho_arquivo = os.path.join(nova_pasta, "declaracao_residencia.docx")
    caminho_arquivo2 = os.path.join(nova_pasta, "declaracao_filiacao.docx")
    doc.save(caminho_arquivo)
    doc2.save(caminho_arquivo2)
    messagebox.showinfo("Sucesso", f"Declaração de residência e filiação de {nome_completo} criadas com sucesso!")

def atualizar_data():
    if var1.get():  # Se estiver marcado
        entry_data.delete(0, tk.END)
        entry_data.insert(0, obter_data_formatada())
    else:
        entry_data.delete(0, tk.END)

def verificar_campos():
    for entry in entries:
        if entry.get().strip() == "":
            messagebox.showwarning("Campo vazio", "Por favor, preencha todos os campos.")
            return False
    return True

def voltar_menu_principal():
    print("Voltando ao Menu Principal")
    root.destroy()
    os.system("python main.py")

def fechar_janela():
    print("fechar")
    root.destroy()

# INTERFACE GRAFICA
root = tk.Tk()
root.title("Gerador de Declarações")
root.config(padx=30, pady=30)

# título H1
tk.Label(root, text="Gerador de Declarações", font=("Arial", 16, "bold")).grid(row=0, column=0, columnspan=2, pady=2)
tk.Label(root, text="Sera gerado as declarações de FILIAÇÃO E RESIDENCIA em .docx", font=("Arial", 10, "bold"),fg="red").grid(row=1, column=0, columnspan=2, pady=0)

labels = ["Nome Completo", "CPF:", "RG", "Estado civil:", 'Rua e número ("Rua X, 123")', "Bairro", "Data de Filiação (DD/MM/AAAA)", 'Data do documento ("5 de maio de 20XX"):']
entries = []

# Preenchendo os campos com Labels e Entry
for i, label in enumerate(labels):
    tk.Label(root, text=label).grid(row=i+2, column=0, padx=5, pady=5)
    if label == "Estado civil:":
        entry_estado_civil = ttk.Combobox(root, values=["Solteiro", "União estável", "Casado"], state="readonly", width=57)
        entry_estado_civil.grid(row=i+2, column=1, padx=5, pady=5)
        entry_estado_civil.current(0)
    else:
        entry = tk.Entry(root, width=60)
        entry.grid(row=i+2, column=1, padx=5, pady=5)
        entries.append(entry)

entry_nome, entry_cpf, entry_rg, entry_rua, entry_bairro, entry_data_filiacao, entry_data = entries

# Campos de entrada
#novo
var1 = tk.BooleanVar()
chk1 = tk.Checkbutton(root, text="Usar data de hoje", variable=var1, command=atualizar_data)
chk1.grid(row=9, column=1, padx=(250,0), pady=5)
#novo
var_sexo = tk.IntVar()
tk.Radiobutton(root, text="Sexo masculino", variable=var_sexo, value=0).grid(row=11, column=1, padx=0, pady=5)
tk.Radiobutton(root, text="Sexo feminino", variable=var_sexo, value=1).grid(row=11, column=0, padx=0, pady=5)

# Botão para gerar documentos
btn_gerar = tk.Button(root, text="Gerar Documentos", bg= "azure3" ,command=gerar_documentos)
btn_gerar.grid(row=12, column=1, padx=10, pady=10, sticky='w')


btn_voltar = tk.Button(root, text="Menu Principal", command=voltar_menu_principal, bg='azure3')
btn_voltar.grid(row=12, column=1, padx=10, pady=20,)


btn_fechar = tk.Button(root, text="Sair", command=fechar_janela, bg='red',fg='white', width=5)
btn_fechar.grid(row=12, column=1, padx=10, pady=10, sticky='e')

root.mainloop()

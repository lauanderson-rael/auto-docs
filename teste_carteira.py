from docx import Document
import os

# Função para substituir texto nas células da tabela
def substituir_texto_tabela(doc, texto_antigo, texto_novo):
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for par in celula.paragraphs:  # Percorrer todos os parágrafos dentro da célula
                    for run in par.runs:  # Percorrer todos os runs dentro do parágrafo
                        if texto_antigo in run.text:
                            run.text = run.text.replace(texto_antigo, texto_novo)

caminho_arquivo = './assets/sociocart2.docx'
doc = Document(caminho_arquivo)

# Definindo as informações para substituição
substituicoes = {
    "NOME_PAI": "RAIMUNDO FRANCISCO BATISTA",
    "NOME_MAE": "MARIA DE LOURDES BATISTA",
    "NOME_RUA": "Rua das Palmeiras, 123",
    "NOME_BAIRRO": "Centro",
    "NATURALIDADE": "São Paulo - SP",
    "DATA_NASC": "01/01/1980",
    "EST_CIVIL": "Casado",
    "NUM_RG_1234567": "12.345.678-9",
    "NUM_CPF1234567": "123.456.789-00",
    "DATA_EMISS": "15/03/2000",
    "N_SIND": "1234",
    "NOME_LOCAL": "Empresa XYZ",
    "NOME_SOCIO": "XLAUANDERSON RAEL SILVA",
}

# Substituir no conteúdo dos parágrafos
for para in doc.paragraphs:
    for run in para.runs:  # Manter a formatação do texto
        for texto_antigo, texto_novo in substituicoes.items():
            if texto_antigo in run.text:
                run.text = run.text.replace(texto_antigo, texto_novo)


# salvar na pasta
nova_pasta = os.path.join("carteiras_socio_geradas")
os.makedirs(nova_pasta, exist_ok=True)
novo_arquivo = os.path.join(nova_pasta, f"{"lauanderson-rael_final"}.docx")
doc.save(novo_arquivo)

print("Carteira de socio criada com sucesso!")

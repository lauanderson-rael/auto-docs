import os
from docx import Document

doc = Document("residencia.docx")
doc2 = Document("filiacao.docx")

texto_informacoes = 'Na falta de documentos próprios, aptos que comprovem minha residência e domicílio, eu, GERIVALDO BARBOSA, Nacionalidade: BRASILEIRO, estado Civil: SOLTEIRO, Profissão: PESCADOR(A), inscrito(a) no Cadastro de Pessoas Físicas (CPF) sob o nº 029.968.823-26, portador(a) da Carteira de Identidade (RG) nº 024627092003-1 declaro ser residente e domiciliado (a) no endereço:  RUA SÃO BRAZ  Bairro: SUBSTAÇÃO'
texto_data = 'Coelho Neto, 28 de abril de 2024'

texto_informacoes2 = "Eu, MARIA HELIELDA SILVA DE FRANÇA, CPF:  083.195.244-07, RG: 067743652018-0, residente no endereço completo: AV COELHO NETO, BAIRRO: CENTRO, declaro ser filiado à Entidade abaixo especificada:"
texto_data2 = "Coelho Neto, 22 de agosto de 2024"

# Dados
nome_completo= 'LAUANDERSON RAEL COSTA GOMES'
cpf = '123.456.789-22'
rg = '1234567892001-0'
sexoF = False
nacionalidade = 'BRASILEIRA' if sexoF else 'BRASILEIRO'
estado_civil = 'UNÃO ESTÁVEL'
rua_e_numero = 'RUA SAO BRAZ, 45'
bairro = 'ANIL'
data = "03 de Fevereiro de 2024"
data_filiacao = "23/01/2018"

# Percorrendo residencia
for para in doc.paragraphs:
    if texto_informacoes in para.text:
        para.text = para.text.replace(texto_informacoes, f'Na falta de documentos próprios, aptos que comprovem minha residência e domicílio, eu, {nome_completo}, Nacionalidade: {nacionalidade}, estado Civil: {estado_civil}, Profissão: PESCADOR(A), inscrito(a) no Cadastro de Pessoas Físicas (CPF) sob o nº {cpf}, portador(a) da Carteira de Identidade (RG) nº {rg}, declaro ser residente e domiciliado (a) no endereço: {rua_e_numero}, Bairro: {bairro}')
    if texto_data in para.text:
        para.text = para.text.replace(texto_data, f'Coelho Neto, {data}')

# Percorrendo filiacao
for para in doc2.paragraphs:
    if texto_informacoes2 in para.text:
        para.text = para.text.replace(texto_informacoes2, f'Eu, {nome_completo}, CPF: {cpf}, RG: {rg}, residente no endereço completo: {rua_e_numero}, Bairro: {bairro}, declaro ser filiado à Entidade abaixo especificada:')
    if "20/08/2015" in para.text:
        novo_texto = para.text.replace("20/08/2015", data_filiacao)
        para.clear()  # Remove o texto original
        para.add_run(novo_texto)
    if texto_data2 in para.text:
        para.text = para.text.replace(texto_data2, f'Coelho Neto, {data}')

# Salvar as alterações
nova_pasta = 'salvos'
os.makedirs(nova_pasta, exist_ok=True)
caminho_arquivo = os.path.join(nova_pasta, f"{cpf}_residencia.docx")
caminho_arquivo2 = os.path.join(nova_pasta, f"{cpf}_filiacao.docx")
doc.save(caminho_arquivo)
doc2.save(caminho_arquivo2)

print(f"Declaração de residência de {nome_completo} criada com sucesso!")
print(f"Declaração de filiação de {nome_completo} criada com sucesso!")

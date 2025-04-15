from docx import Document
import os
import tkinter as tk
from tkinter import messagebox

# Função para substituir texto nas células da tabela sem perder a formatação
def substituir_texto_tabela(doc, texto_antigo, texto_novo):
    # Loop através de todas as tabelas no documento
    for tabela in doc.tables:
        # Loop através de todas as linhas e células da tabela
        for linha in tabela.rows:
            for celula in linha.cells:
                for par in celula.paragraphs:  # Percorrer todos os parágrafos dentro da célula
                    for run in par.runs:  # Percorrer todos os runs dentro do parágrafo
                        if texto_antigo in run.text:
                            run.text = run.text.replace(texto_antigo, texto_novo)

# Função para gerar a carteira
def gerar_carteira():
    # Pegando os dados dos campos de entrada
    nome_pai = entry_nome_pai.get()
    nome_mae = entry_nome_mae.get()
    nome_rua = entry_nome_rua.get()
    nome_bairro = entry_nome_bairro.get()
    naturalidade = entry_naturalidade.get()
    data_nasc = entry_data_nasc.get()
    estado_civil = entry_estado_civil.get()
    num_rg = entry_num_rg.get()
    num_cpf = entry_num_cpf.get()
    data_emissao = entry_data_emissao.get()
    numero_sind = entry_numero_sind.get()
    nome_local = entry_nome_local.get()
    nome_socios = entry_nome_socios.get()

    # Caminho para o arquivo Word
    caminho_arquivo = 'sociocart2.docx'

    # Verifique se o arquivo de entrada existe
    if not os.path.exists(caminho_arquivo):
        messagebox.showerror("Erro", "O arquivo de entrada 'sociocart2.docx' não foi encontrado!")
        return

    # Carregar o documento
    doc = Document(caminho_arquivo)

    # Definindo as informações para substituição
    substituicoes = {
        "NOME_PAI": nome_pai,
        "NOME_MAE": nome_mae,
        "NOME_RUA": nome_rua,
        "NOME_BAIRRO": nome_bairro,
        "NATURALIDADE": naturalidade,
        "DATA_NASC": data_nasc,
        "EST_CIVIL": estado_civil,
        "NUM_RG1234567": num_rg,
        "NUM_CPF1234567": num_cpf,
        "DATA_EMISS": data_emissao,
        "N_SIND": numero_sind,
        "NOME_LOCAL": nome_local,
        "NOME_SOCIO": nome_socios,
    }

    # Substituir no conteúdo dos parágrafos
    for para in doc.paragraphs:
        for run in para.runs:  # Manter a formatação do texto
            for texto_antigo, texto_novo in substituicoes.items():
                if texto_antigo in run.text:
                    run.text = run.text.replace(texto_antigo, texto_novo)

    # Criar nova pasta onde o arquivo será salvo, se não existir
    nova_pasta = os.path.join("carteiras_socio_geradas")
    os.makedirs(nova_pasta, exist_ok=True)

    # Definir o caminho completo para o novo arquivo
    novo_arquivo = os.path.join(nova_pasta, "lauanderson-rael_final.docx")

    # Salvar o documento com as substituições
    doc.save(novo_arquivo)

    # Exibir mensagem de sucesso
    messagebox.showinfo("Sucesso", "A carteira foi criada com sucesso!")

# Criando a janela principal
root = tk.Tk()
root.title("Gerador de Carteira de Sócio")

# Definindo o layout da interface gráfica
tk.Label(root, text="Nome do Pai:").grid(row=0, column=0, padx=10, pady=5)
entry_nome_pai = tk.Entry(root)
entry_nome_pai.grid(row=0, column=1, padx=10, pady=5)

tk.Label(root, text="Nome da Mãe:").grid(row=1, column=0, padx=10, pady=5)
entry_nome_mae = tk.Entry(root)
entry_nome_mae.grid(row=1, column=1, padx=10, pady=5)

tk.Label(root, text="Nome da Rua:").grid(row=2, column=0, padx=10, pady=5)
entry_nome_rua = tk.Entry(root)
entry_nome_rua.grid(row=2, column=1, padx=10, pady=5)

tk.Label(root, text="Nome do Bairro:").grid(row=3, column=0, padx=10, pady=5)
entry_nome_bairro = tk.Entry(root)
entry_nome_bairro.grid(row=3, column=1, padx=10, pady=5)

tk.Label(root, text="Naturalidade:").grid(row=4, column=0, padx=10, pady=5)
entry_naturalidade = tk.Entry(root)
entry_naturalidade.grid(row=4, column=1, padx=10, pady=5)

tk.Label(root, text="Data de Nascimento:").grid(row=5, column=0, padx=10, pady=5)
entry_data_nasc = tk.Entry(root)
entry_data_nasc.grid(row=5, column=1, padx=10, pady=5)

tk.Label(root, text="Estado Civil:").grid(row=6, column=0, padx=10, pady=5)
entry_estado_civil = tk.Entry(root)
entry_estado_civil.grid(row=6, column=1, padx=10, pady=5)

tk.Label(root, text="Número do RG:").grid(row=7, column=0, padx=10, pady=5)
entry_num_rg = tk.Entry(root)
entry_num_rg.grid(row=7, column=1, padx=10, pady=5)

tk.Label(root, text="Número do CPF:").grid(row=8, column=0, padx=10, pady=5)
entry_num_cpf = tk.Entry(root)
entry_num_cpf.grid(row=8, column=1, padx=10, pady=5)

tk.Label(root, text="Data de Emissão:").grid(row=9, column=0, padx=10, pady=5)
entry_data_emissao = tk.Entry(root)
entry_data_emissao.grid(row=9, column=1, padx=10, pady=5)

tk.Label(root, text="Número do Sindicado:").grid(row=10, column=0, padx=10, pady=5)
entry_numero_sind = tk.Entry(root)
entry_numero_sind.grid(row=10, column=1, padx=10, pady=5)

tk.Label(root, text="Nome do Local de Trabalho:").grid(row=11, column=0, padx=10, pady=5)
entry_nome_local = tk.Entry(root)
entry_nome_local.grid(row=11, column=1, padx=10, pady=5)

tk.Label(root, text="Nome do Sócio:").grid(row=12, column=0, padx=10, pady=5)
entry_nome_socios = tk.Entry(root)
entry_nome_socios.grid(row=12, column=1, padx=10, pady=5)

# Botão para gerar a carteira
btn_gerar = tk.Button(root, text="Gerar Carteira", command=gerar_carteira)
btn_gerar.grid(row=13, column=0, columnspan=2, pady=20)

# Rodar a interface gráfica
root.mainloop()
